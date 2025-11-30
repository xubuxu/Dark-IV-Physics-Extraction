from docx import Document
from docx.shared import Inches
import os
import pandas as pd

class WordReportBuilder:
    def __init__(self, template_path):
        if os.path.exists(template_path):
            self.doc = Document(template_path)
        else:
            print(f"Template not found: {template_path}. Creating new document.")
            self.doc = Document()
            self.doc.add_heading('Dark IV Analysis Report', 0)
            
    def add_header(self, text, level=1):
        self.doc.add_heading(text, level)
        
    def add_summary_table(self, df):
        """Adds a summary table to the document."""
        self.doc.add_heading('Summary Table', level=2)
        
        # Limit columns for Word table readability
        cols_to_show = ['filename', 'J0', 'n', 'Rs', 'Rsh']
        # Filter existing
        cols_to_show = [c for c in cols_to_show if c in df.columns]
        
        t = self.doc.add_table(rows=1, cols=len(cols_to_show))
        t.style = 'Table Grid'
        
        # Header
        hdr_cells = t.rows[0].cells
        for i, col in enumerate(cols_to_show):
            hdr_cells[i].text = str(col)
            
        # Rows
        for index, row in df.iterrows():
            row_cells = t.add_row().cells
            for i, col in enumerate(cols_to_show):
                val = row[col]
                if isinstance(val, float):
                    row_cells[i].text = f"{val:.2e}" if abs(val) < 1e-3 or abs(val) > 1e3 else f"{val:.4f}"
                else:
                    row_cells[i].text = str(val)
                
    def add_sample_analysis(self, sample_name, params, plot_streams):
        """
        Adds analysis section for a single sample.
        
        Args:
            sample_name: Name of the sample
            params: Dict of extracted parameters
            plot_streams: Either a single BytesIO stream or dict of streams {'jv': stream, 'n': stream, 'rdiff': stream}
        """
        self.doc.add_heading(f"Sample: {sample_name}", level=2)
        
        # Add parameters text
        p_text = "Extracted Parameters:\n"
        for k, v in params.items():
            if k == 'filename': continue
            if isinstance(v, float):
                val_str = f"{v:.2e}" if (abs(v) < 1e-3 or abs(v) > 1e3) else f"{v:.4f}"
            else:
                val_str = str(v)
            p_text += f"{k}: {val_str}\n"
            
        self.doc.add_paragraph(p_text)
        
        # Add images
        if isinstance(plot_streams, dict):
            # Multiple plots
            plot_titles = {
                'jv': 'Semi-log Dark J-V Curve',
                'n': 'Local Ideality Factor n(V)',
                'rdiff': 'Differential Resistance Rdiff(V)'
            }
            
            for key in ['jv', 'n', 'rdiff']:
                if key in plot_streams and plot_streams[key]:
                    self.doc.add_heading(plot_titles.get(key, key), level=3)
                    self.doc.add_picture(plot_streams[key], width=Inches(6.0))
                    self.doc.add_paragraph()  # Add spacing
        else:
            # Single plot (backward compatibility)
            if plot_streams:
                self.doc.add_picture(plot_streams, width=Inches(6.0))
        
    def save(self, filepath):
        self.doc.save(filepath)
